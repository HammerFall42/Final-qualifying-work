from docx import Document
from docx.shared import Pt

class Docformer:
    def checkDuplicates(self,  strings):
        i = 0
        while i < len(strings):
            j = 0
            while j < len(strings):
                if i == j:
                    j += 1
                    continue
                if strings[j][0] == strings[i][0]:
                    del(strings[j])
                    j -= 1
                j += 1
            i += 1

        return strings

    def createDoc(self, strings, filename):
        if strings == []:
            return
        strings = self.checkDuplicates(strings)
        document = Document()
        par_created = False
        para = None
        list_goes = False
        list_count = 1
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        for item in strings:
            if item[1] == "h":
                document.add_heading(item[0])
                par_created = False
                list_goes = False
            elif item[1] == "p":
                list_goes = False
                if not par_created:
                    para = document.add_paragraph(item[0], style=document.styles['Normal'])
                    par_created = True
                else:
                    para.add_run(item[0])
            else:
                if list_goes:
                    list_count += 1
                else:
                    list_count = 1
                    list_goes = True
                document.add_paragraph(str(list_count) + ". " + item[0], style=document.styles['Normal'])
                par_created = False
        document.save(filename + ".docx")
